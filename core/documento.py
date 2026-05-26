from docx import Document
import os
from docx2pdf import convert
import platform
import subprocess
from docx.shared import Inches

def gerar_contrato_word(caminho_modelo, caminho_saida, dados):
    """
    Lê o modelo Word, substitui as tags de texto e injeta a imagem da 
    assinatura no lugar da tag {{ASSINATURA}}.
    """
    doc = Document(caminho_modelo)
    
    # 1. Substituição nos parágrafos comuns (Texto livre, assinaturas, etc.)
    for paragrafo in doc.paragraphs:
        for tag, valor in dados.items():
            if tag in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(tag, str(valor))
                
    # 2. Substituição dentro das Tabelas do documento
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for tag, valor in dados.items():
                    if tag in celula.text:
                        celula.text = celula.text.replace(tag, str(valor))

    # 3. MÁGICA DA ASSINATURA: Injeta a imagem exatamente no parágrafo da tag
    caminho_assinatura = "templates/assinatura.png"
    
    for paragrafo in doc.paragraphs:
        if "{{ASSINATURA}}" in paragrafo.text:
            # Limpa o texto da tag, deixando o parágrafo vazio, mas mantendo a sua posição
            paragrafo.text = paragrafo.text.replace("{{ASSINATURA}}", "")
            
            if os.path.exists(caminho_assinatura):
                # Adiciona a imagem direto no parágrafo limpo
                run = paragrafo.add_run()
                run.add_picture(caminho_assinatura, width=Inches(2.0))
            break
                
                        
    # 4. Salva o documento preenchido
    doc.save(caminho_saida)

# Converter o documento em word para pdf
def converter_docx_para_pdf(caminho_docx, caminho_pdf):
    """
    Função inteligente que converte Word para PDF tanto no Windows local
    quanto no servidor Linux da nuvem do Streamlit.
    """
    try:
        # Detecta qual é o sistema operacional atual (Windows, Linux, etc.)
        sistema = platform.system()
        
        if sistema == "Windows":
            print("Ambiente Local (Windows) detectado. Usando docx2pdf...")
            convert(caminho_docx, caminho_pdf)
            return True
            
        elif sistema == "Linux":
            print("Ambiente de Nuvem (Linux) detectado. Usando LibreOffice...")
            # Pega a pasta onde o arquivo está guardado
            pasta_saida = os.path.dirname(caminho_pdf)
            
            # Comando mágico do Linux que pede para o LibreOffice converter o arquivo de forma invisível
            comando = [
                "libreoffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", pasta_saida,
                caminho_docx
            ]
            
            # Executa o comando no servidor Linux
            subprocess.run(comando, check=True)
            return True
            
        else:
            print(f"Sistema operacional {sistema} não suportado para conversão automática.")
            return False
            
    except Exception as e:
        print(f"Erro na conversão para PDF: {e}")
        return False